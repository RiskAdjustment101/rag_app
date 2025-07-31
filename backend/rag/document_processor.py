"""
Document Processing Pipeline for RAG
Handles PDF, DOCX, and TXT file parsing and chunking
"""
import os
import re
import uuid
from typing import List, Dict, Any, Tuple
from io import BytesIO
import logging
import aiofiles
import tiktoken

# Document parsing libraries
import pypdf
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process documents for RAG ingestion"""
    
    def __init__(self):
        # Initialize tokenizer for chunk size estimation
        self.tokenizer = tiktoken.get_encoding("cl100k_base")  # GPT-4 tokenizer
        self.max_chunk_size = 1000  # tokens
        self.chunk_overlap = 200    # tokens
    
    async def process_document(
        self, 
        file_content: bytes, 
        filename: str, 
        user_id: str
    ) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Process a document and return chunks with metadata
        Returns: (document_id, chunks_with_metadata)
        """
        try:
            # Generate unique document ID
            document_id = f"{user_id}_{uuid.uuid4().hex}"
            
            # Extract text based on file type
            file_extension = filename.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                text = await self._extract_pdf_text(file_content)
            elif file_extension in ['docx', 'doc']:
                text = await self._extract_docx_text(file_content)
            elif file_extension == 'txt':
                text = await self._extract_txt_text(file_content)
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            if not text or len(text.strip()) < 50:
                raise ValueError("Document appears to be empty or too short")
            
            # Clean and chunk the text
            cleaned_text = self._clean_text(text)
            chunks = self._chunk_text(cleaned_text)
            
            # Create metadata for each chunk
            chunks_with_metadata = []
            for i, chunk in enumerate(chunks):
                metadata = {
                    "filename": filename,
                    "file_type": file_extension,
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "char_count": len(chunk),
                    "token_count": len(self.tokenizer.encode(chunk)),
                    "document_id": document_id
                }
                chunks_with_metadata.append({
                    "text": chunk,
                    "metadata": metadata
                })
            
            logger.info(f"Processed document {filename}: {len(chunks)} chunks extracted")
            return document_id, chunks_with_metadata
            
        except Exception as e:
            logger.error(f"Failed to process document {filename}: {e}")
            raise
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_content = []
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content:
                raise ValueError("No readable text found in PDF")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Failed to process PDF: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            if not text_content:
                raise ValueError("No readable text found in DOCX")
            
            return "\n\n".join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")
    
    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    if text.strip():
                        return text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise ValueError(f"Failed to process TXT: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content"""
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove or replace special characters
        text = text.replace('\r', '\n')
        text = text.replace('\t', ' ')
        
        # Remove excessive line breaks but preserve paragraph structure
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
    
    def _chunk_text(self, text: str) -> List[str]:
        """Split text into chunks with overlap"""
        # Split text into sentences for better chunk boundaries
        sentences = self._split_into_sentences(text)
        
        chunks = []
        current_chunk = []
        current_tokens = 0
        
        for sentence in sentences:
            sentence_tokens = len(self.tokenizer.encode(sentence))
            
            # If adding this sentence would exceed the limit, finish current chunk
            if current_tokens + sentence_tokens > self.max_chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                
                # Start new chunk with overlap
                overlap_chunk = []
                overlap_tokens = 0
                
                # Add sentences from the end of current chunk for overlap
                for prev_sentence in reversed(current_chunk):
                    prev_tokens = len(self.tokenizer.encode(prev_sentence))
                    if overlap_tokens + prev_tokens <= self.chunk_overlap:
                        overlap_chunk.insert(0, prev_sentence)
                        overlap_tokens += prev_tokens
                    else:
                        break
                
                current_chunk = overlap_chunk + [sentence]
                current_tokens = overlap_tokens + sentence_tokens
            else:
                current_chunk.append(sentence)
                current_tokens += sentence_tokens
        
        # Add the last chunk if it has content
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        # Simple sentence splitting - could be improved with NLTK
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Clean up sentences and filter out very short ones
        cleaned_sentences = []
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) > 10:  # Minimum sentence length
                cleaned_sentences.append(sentence)
        
        return cleaned_sentences

# Global document processor instance
document_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get or create global document processor instance"""
    global document_processor
    if document_processor is None:
        document_processor = DocumentProcessor()
    return document_processor